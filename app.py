from flask import Flask, render_template, request, jsonify, send_file, session
import os
import uuid
import fitz  # PyMuPDF
import base64
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches

app = Flask(__name__)
app.secret_key = 'pdf_signature_secret_key_2024'

# Configurações
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
ALLOWED_EXTENSIONS = {'pdf'}

# Criar pastas se não existirem
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def convert_coordinates_to_pdf(marker_x, marker_y, page_height):
    """Converte coordenadas da marcação para coordenadas PDF"""
    # Conversão de coordenadas (2x zoom para 1x PDF)
    pdf_x = marker_x / 2.0
    pdf_y = page_height - (marker_y / 2.0)
    
    # Ajuste para baseline
    baseline_offset = 15.0
    pdf_y = pdf_y - baseline_offset
    
    return pdf_x, pdf_y

def validate_signature_position(pdf_x, pdf_y, page_width, page_height):
    """Valida e ajusta posição da assinatura dentro dos limites da página"""
    signature_width = 300
    signature_height = 80
    margin = 20
    
    final_x = max(margin, min(pdf_x, page_width - signature_width - margin))
    final_y = max(margin, min(pdf_y, page_height - signature_height - margin))
    
    return final_x, final_y

def render_pdf_page_as_image(pdf_path, page_num, zoom=2.0):
    """Renderiza uma página específica do PDF como imagem"""
    try:
        doc = fitz.open(pdf_path)
        if page_num >= len(doc):
            return None
        
        page = doc.load_page(page_num)
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        img_data = pix.tobytes("png")
        img_base64 = base64.b64encode(img_data).decode()
        
        result = {
            'page_num': page_num + 1,
            'image': f"data:image/png;base64,{img_base64}",
            'width': pix.width,
            'height': pix.height
        }
        
        doc.close()
        return result
    except Exception as e:
        print(f"Erro ao renderizar página {page_num + 1}: {e}")
        return None

def get_pdf_pages_as_images(pdf_path, max_pages=10):
    """Converte páginas do PDF em imagens para visualização"""
    try:
        doc = fitz.open(pdf_path)
        images = []
        
        for page_num in range(min(len(doc), max_pages)):
            result = render_pdf_page_as_image(pdf_path, page_num)
            if result:
                images.append(result)
        
        doc.close()
        return images
    except Exception as e:
        print(f"Erro ao processar PDF: {e}")
        return []

@app.route('/')
def index():
    """Página principal com todas as ferramentas"""
    return render_template('main.html')

@app.route('/sign')
def sign_page():
    """Página de assinatura digital"""
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': 'Nenhum arquivo selecionado'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': 'Nenhum arquivo selecionado'})
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            unique_filename = f"{uuid.uuid4()}_{filename}"
            filepath = os.path.join(UPLOAD_FOLDER, unique_filename)
            file.save(filepath)
            
            # Gerar imagens das páginas
            images = get_pdf_pages_as_images(filepath)
            
            if not images:
                return jsonify({'success': False, 'message': 'Erro ao processar PDF'})
            
            # Salvar informações na sessão
            session['current_pdf'] = filepath
            session['pdf_filename'] = filename
            
            return jsonify({
                'success': True, 
                'message': 'PDF carregado com sucesso',
                'images': images,
                'filename': filename
            })
        else:
            return jsonify({'success': False, 'message': 'Tipo de arquivo não permitido'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro: {str(e)}'})

@app.route('/api/signatures', methods=['GET'])
def get_signatures():
    """API para gerenciar assinaturas eletrônicas"""
    try:
        # Assinaturas pré-definidas
        signatures = [
            {
                'id': 1,
                'name': 'João Silva',
                'position': 'Diretor',
                'signature_type': 'digital',
                'created_at': '2024-01-15',
                'status': 'active'
            },
            {
                'id': 2,
                'name': 'Maria Santos',
                'position': 'Gerente',
                'signature_type': 'digital',
                'created_at': '2024-01-20',
                'status': 'active'
            },
            {
                'id': 3,
                'name': 'Pedro Costa',
                'position': 'Coordenador',
                'signature_type': 'digital',
                'created_at': '2024-01-25',
                'status': 'active'
            }
        ]
        
        return jsonify({
            'success': True,
            'signatures': signatures,
            'total': len(signatures)
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro: {str(e)}'})

@app.route('/api/signatures/auto-load', methods=['POST'])
def auto_load_signature():
    """API para carregar assinatura automaticamente"""
    try:
        data = request.json
        pdf_path = session.get('current_pdf')
        
        if not pdf_path or not os.path.exists(pdf_path):
            return jsonify({'success': False, 'message': 'PDF não encontrado'})
        
        # Dados da assinatura
        signer_name = data.get('signer_name', 'Assinatura Digital').strip()
        signer_position = data.get('signer_position', 'Assinado Digitalmente').strip()
        page_num = int(data.get('page_num', 1)) - 1
        
        # Obter coordenadas
        doc = fitz.open(pdf_path)
        if page_num >= len(doc):
            return jsonify({'success': False, 'message': 'Página inválida'})
        
        page = doc.load_page(page_num)
        page_rect = page.rect
        page_width = page_rect.width
        page_height = page_rect.height
        
        # Coordenadas da marcação
        marker_x = float(data.get('marker_x', page_width / 2))
        marker_y = float(data.get('marker_y', page_height / 2))
        
        # Conversão e validação de coordenadas
        pdf_x, pdf_y = convert_coordinates_to_pdf(marker_x, marker_y, page_height)
        final_x, final_y = validate_signature_position(pdf_x, pdf_y, page_width, page_height)
        
        doc.close()
        
        # Retornar PDF com coordenadas
        with open(pdf_path, 'rb') as file:
            pdf_bytes = file.read()
        
        return jsonify({
            'success': True,
            'pdf_data': base64.b64encode(pdf_bytes).decode('utf-8'),
            'coordinates': {
                'x': final_x,
                'y': final_y,
                'page': page_num,
                'page_width': page_width,
                'page_height': page_height,
                'signature_width': signature_width,
                'signature_height': signature_height,
                'marker_x': marker_x,
                'marker_y': marker_y
            },
            'signature_info': {
                'name': signer_name,
                'position': signer_position,
                'auto_loaded': True
            }
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro: {str(e)}'})

@app.route('/add_signature_frontend', methods=['POST'])
def add_signature_frontend():
    """API principal para assinatura com posicionamento preciso"""
    try:
        data = request.json
        pdf_path = session.get('current_pdf')
        
        if not pdf_path or not os.path.exists(pdf_path):
            return jsonify({'success': False, 'message': 'PDF não encontrado'})
        
        # Dados da assinatura
        signer_name = data.get('signer_name', '').strip()
        signer_position = data.get('signer_position', '').strip()
        page_num = int(data.get('page_num', 1)) - 1
        
        if not signer_name:
            return jsonify({'success': False, 'message': 'Nome do signatário é obrigatório'})
        
        # Obter coordenadas precisas
        doc = fitz.open(pdf_path)
        if page_num >= len(doc):
            return jsonify({'success': False, 'message': 'Página inválida'})
        
        page = doc.load_page(page_num)
        page_rect = page.rect
        page_width = page_rect.width
        page_height = page_rect.height
        
        # Coordenadas da marcação
        marker_x = float(data.get('x', page_width / 2))
        marker_y = float(data.get('y', page_height / 2))
        
        # Conversão e validação de coordenadas
        pdf_x, pdf_y = convert_coordinates_to_pdf(marker_x, marker_y, page_height)
        final_x, final_y = validate_signature_position(pdf_x, pdf_y, page_width, page_height)
        
        doc.close()
        
        # Retornar PDF com coordenadas
        with open(pdf_path, 'rb') as file:
            pdf_bytes = file.read()
        
        return jsonify({
            'success': True,
            'pdf_data': base64.b64encode(pdf_bytes).decode('utf-8'),
            'coordinates': {
                'x': final_x,
                'y': final_y,
                'page': page_num,
                'page_width': page_width,
                'page_height': page_height,
                'signature_width': signature_width,
                'signature_height': signature_height
            }
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro: {str(e)}'})

@app.route('/download/<filename>')
def download_file(filename):
    try:
        file_path = os.path.join(OUTPUT_FOLDER, filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            return jsonify({'error': 'Arquivo não encontrado'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/split')
def split_page():
    """Página para dividir PDF"""
    return render_template('split.html')

@app.route('/merge')
def merge_page():
    """Página para juntar PDFs"""
    return render_template('merge.html')

@app.route('/convert')
def convert_page():
    """Página para converter PDF para Word"""
    return render_template('convert.html')

@app.route('/convert_pdf_to_word', methods=['POST'])
def convert_pdf_to_word():
    """Converter PDF para Word"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': 'Nenhum arquivo enviado'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': 'Nenhum arquivo selecionado'})
        
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({'success': False, 'message': 'Por favor, envie um arquivo PDF'})
        
        # Salvar arquivo temporário
        filename = secure_filename(file.filename)
        temp_filename = f"{uuid.uuid4()}_{filename}"
        temp_path = os.path.join(UPLOAD_FOLDER, temp_filename)
        file.save(temp_path)
        
        try:
            # Converter PDF para Word
            pdf_doc = fitz.open(temp_path)
            doc = Document()
            
            # Processar cada página
            for page_num in range(len(pdf_doc)):
                page = pdf_doc.load_page(page_num)
                text = page.get_text()
                
                if text.strip():
                    paragraph = doc.add_paragraph()
                    lines = text.split('\n')
                    for line in lines:
                        if line.strip():
                            run = paragraph.add_run(line.strip() + '\n')
                            if line.strip().isupper() or len(line.strip()) < 50:
                                run.bold = True
                
                # Quebra de página
                if page_num < len(pdf_doc) - 1:
                    doc.add_page_break()
                
                # Incluir imagens
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(pdf_doc, xref)
                        
                        if pix.n - pix.alpha < 4:
                            img_filename = f"temp_img_{page_num}_{img_index}.png"
                            img_path = os.path.join(UPLOAD_FOLDER, img_filename)
                            pix.save(img_path)
                            doc.add_picture(img_path, width=Inches(4))
                            os.remove(img_path)
                        
                        pix = None
                    except Exception as e:
                        continue
            
            pdf_doc.close()
            
            # Salvar documento Word
            word_filename = f"converted_{uuid.uuid4().hex[:8]}.docx"
            word_path = os.path.join(OUTPUT_FOLDER, word_filename)
            doc.save(word_path)
            
            return jsonify({
                'success': True,
                'filename': word_filename,
                'message': f'PDF convertido para Word com sucesso!'
            })
            
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro: {str(e)}'})

@app.route('/download_word/<filename>')
def download_word_file(filename):
    """Download de arquivo Word convertido"""
    try:
        file_path = os.path.join(OUTPUT_FOLDER, filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            return jsonify({'error': 'Arquivo não encontrado'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/merge_pdf', methods=['POST'])
def merge_pdf():
    """Juntar múltiplos PDFs em um único arquivo"""
    try:
        if 'files' not in request.files:
            return jsonify({'success': False, 'message': 'Nenhum arquivo enviado'})
        
        files = request.files.getlist('files')
        if not files or all(file.filename == '' for file in files):
            return jsonify({'success': False, 'message': 'Nenhum arquivo selecionado'})
        
        # Filtrar apenas PDFs
        pdf_files = [file for file in files if file.filename.lower().endswith('.pdf')]
        
        if len(pdf_files) < 2:
            return jsonify({'success': False, 'message': 'É necessário pelo menos 2 arquivos PDF para juntar'})
        
        # Criar documento PDF final
        merged_doc = fitz.open()
        
        # Processar cada arquivo PDF
        for file in pdf_files:
            # Salvar arquivo temporário
            filename = secure_filename(file.filename)
            temp_filename = f"{uuid.uuid4()}_{filename}"
            temp_path = os.path.join(UPLOAD_FOLDER, temp_filename)
            file.save(temp_path)
            
            try:
                # Abrir PDF e inserir páginas
                doc = fitz.open(temp_path)
                merged_doc.insert_pdf(doc)
                doc.close()
            finally:
                # Remover arquivo temporário
                if os.path.exists(temp_path):
                    os.remove(temp_path)
        
        # Salvar PDF final
        final_filename = f"merged_{uuid.uuid4().hex[:8]}.pdf"
        final_path = os.path.join(OUTPUT_FOLDER, final_filename)
        merged_doc.save(final_path)
        merged_doc.close()
        
        return jsonify({
            'success': True,
            'filename': final_filename,
            'message': f'{len(pdf_files)} PDFs juntados com sucesso'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro: {str(e)}'})

@app.route('/download_merge/<filename>')
def download_merge_file(filename):
    """Download de arquivo juntado"""
    try:
        file_path = os.path.join(OUTPUT_FOLDER, filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            return jsonify({'error': 'Arquivo não encontrado'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/split_pdf', methods=['POST'])
def split_pdf():
    """Carregar PDF para divisão"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': 'Nenhum arquivo enviado'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': 'Nenhum arquivo selecionado'})
        
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({'success': False, 'message': 'Por favor, envie um arquivo PDF'})
        
        # Salvar arquivo
        filename = secure_filename(file.filename)
        unique_filename = f"{uuid.uuid4()}_{filename}"
        file_path = os.path.join(UPLOAD_FOLDER, unique_filename)
        file.save(file_path)
        
        # Armazenar na sessão
        session['current_pdf'] = file_path
        
        # Abrir PDF e extrair páginas
        doc = fitz.open(file_path)
        pages = []
        
        for page_num in range(len(doc)):
            result = render_pdf_page_as_image(file_path, page_num)
            if result:
                pages.append({
                    'page_num': result['page_num'],
                    'image': result['image']
                })
        
        doc.close()
        
        return jsonify({
            'success': True,
            'filename': filename,
            'pages': pages,
            'total_pages': len(pages)
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro: {str(e)}'})

@app.route('/split_pdf_process', methods=['POST'])
def split_pdf_process():
    """Processar divisão do PDF"""
    try:
        data = request.json
        pdf_path = session.get('current_pdf')
        
        if not pdf_path or not os.path.exists(pdf_path):
            return jsonify({'success': False, 'message': 'PDF não encontrado'})
        
        method = data.get('method')
        value = data.get('value')
        
        # Abrir PDF
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        
        split_files = []
        
        # Apenas método de páginas selecionadas
        if method == 'selected':
            # Dividir apenas páginas selecionadas (não excluídas)
            excluded_pages = set(data.get('excluded_pages', []))
            available_pages = [i+1 for i in range(total_pages) if (i+1) not in excluded_pages]
            
            if not available_pages:
                return jsonify({'success': False, 'message': 'Todas as páginas foram excluídas'})
            
            # Criar um PDF com todas as páginas disponíveis
            new_doc = fitz.open()
            for page_num in available_pages:
                new_doc.insert_pdf(doc, from_page=page_num-1, to_page=page_num-1)
            
            filename = f"split_selected_{uuid.uuid4().hex[:8]}.pdf"
            file_path = os.path.join(OUTPUT_FOLDER, filename)
            new_doc.save(file_path)
            new_doc.close()
            split_files.append(filename)
        
        doc.close()
        
        return jsonify({
            'success': True,
            'files': split_files,
            'message': f'PDF dividido em {len(split_files)} arquivos'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro: {str(e)}'})

@app.route('/download_split/<filename>')
def download_split_file(filename):
    """Download de arquivo dividido"""
    try:
        file_path = os.path.join(OUTPUT_FOLDER, filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            return jsonify({'error': 'Arquivo não encontrado'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)